# -*- coding: utf-8 -*-
"""Builds Lendora_Speaking_Script.docx — full rehearsal script for all 39 slides."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from slides_data import SLIDES

FOREST = RGBColor(0x0A, 0x3D, 0x2E)
COPPER = RGBColor(0xC8, 0x79, 0x41)
INK = RGBColor(0x22, 0x26, 0x24)
INK_LT = RGBColor(0x5A, 0x60, 0x5C)

doc = Document()

# base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = INK


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): hex_color})
    tcPr.append(shd)


# ---- cover page ----
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("LENDORA")
r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = FOREST; r.font.name = 'Georgia'

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Full Speaking Script — University Project Defense")
r.font.size = Pt(16); r.font.italic = True; r.font.color.rgb = COPPER

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run(
    "39 slides  |  Target length: 15–20 minutes  |  ~30–60 seconds of speech per slide\n"
    "Read this script aloud once or twice before the defense, then speak from memory using the\n"
    "bullet recap on each slide as your anchor — don't read word-for-word in the room."
)
r.font.size = Pt(11); r.font.color.rgb = INK_LT

doc.add_page_break()

# ---- table of contents (simple list) ----
h = doc.add_heading("Slide Overview", level=1)
h.runs[0].font.color.rgb = FOREST
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "#", "Slide Title", "Type"
for c in hdr:
    shade_cell(c, "0A3D2E")
    for p in c.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.bold = True

KIND_LABEL = {
    'title': 'Title', 'content': 'Content', 'screenshot': 'Screenshot',
    'diagram_arch': 'Diagram', 'diagram_workflow': 'Diagram', 'diagram_db': 'Diagram',
    'thanks': 'Closing',
}
for s in SLIDES:
    row = table.add_row().cells
    row[0].text = str(s['num'])
    row[1].text = s['title']
    row[2].text = KIND_LABEL[s['kind']]

doc.add_page_break()

# ---- per-slide script pages ----
for s in SLIDES:
    h = doc.add_heading(f"Slide {s['num']} — {s['title']}", level=1)
    h.runs[0].font.color.rgb = FOREST

    if s['kind'] in ('title', 'thanks'):
        recap = [s['subtitle']] + s.get('meta', [])
    elif s['kind'] == 'screenshot':
        recap = [
            f"What it does: {s['what']}",
            f"Who uses it: {s['who']}",
            f"Key actions: {s['actions']}",
            f"Leads to: {s['connects']}",
        ]
    elif s['kind'] == 'diagram_workflow':
        recap = ["Flow: " + " → ".join(st.replace(chr(10), ' ') for st in s['steps'])]
    else:
        recap = s.get('bullets', [])

    p = doc.add_paragraph()
    r = p.add_run("On-slide content:")
    r.bold = True; r.font.color.rgb = COPPER
    for line in recap:
        bp = doc.add_paragraph(style='List Bullet')
        bp.add_run(line)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Speaking script (~30–60 sec):")
    r2.bold = True; r2.font.color.rgb = COPPER
    sp = doc.add_paragraph(s['script'])
    sp.paragraph_format.space_after = Pt(10)

    if s.get('transition'):
        tp = doc.add_paragraph()
        r3 = tp.add_run("➡ Transition to next slide: ")
        r3.bold = True; r3.font.color.rgb = FOREST
        tp.add_run(s['transition']).italic = True

    if s.get('shots'):
        sh = doc.add_paragraph()
        r4 = sh.add_run("Screenshot(s) to insert: ")
        r4.bold = True; r4.font.color.rgb = INK_LT
        sh.add_run(" | ".join(s['shots'])).italic = True

    if s['num'] != 39:
        doc.add_page_break()

doc.save("Lendora_Speaking_Script.docx")
print("Saved Lendora_Speaking_Script.docx")
